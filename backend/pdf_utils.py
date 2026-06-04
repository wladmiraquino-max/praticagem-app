"""Utilitários de extração de texto com múltiplos parsers como fallback."""
import os
import tempfile


def extract_text_from_pdf(raw_bytes: bytes, max_pages: int = 50) -> str:
    """Tenta extrair texto de um PDF usando pdfplumber e pypdf como fallback."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(raw_bytes)
        tmp_path = tmp.name

    text = ""
    try:
        # Tentativa 1: pdfplumber (melhor para PDFs com texto nativo)
        try:
            import pdfplumber
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages[:max_pages]:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
        except Exception:
            pass

        # Tentativa 2: pypdf como fallback
        if not text.strip():
            try:
                from pypdf import PdfReader
                reader = PdfReader(tmp_path)
                for page in reader.pages[:max_pages]:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            except Exception:
                pass

    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    return text


def extract_text_from_docx(raw_bytes: bytes) -> str:
    """Extrai texto de arquivos .docx."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(raw_bytes)
        tmp_path = tmp.name
    text = ""
    try:
        from docx import Document
        doc = Document(tmp_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " | "
                text += "\n"
    except Exception:
        pass
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
    return text


def extract_text(raw_bytes: bytes, filename: str, max_pages: int = 50) -> tuple[str, str | None]:
    """
    Extrai texto de PDF, DOCX ou TXT.
    Retorna (content, error_message). Se error_message não for None, exiba ao usuário.
    """
    fname = filename.lower()

    if fname.endswith(".pdf"):
        text = extract_text_from_pdf(raw_bytes, max_pages)
        if not text.strip():
            return "", (
                "Não foi possível extrair texto deste PDF. "
                "Ele pode ser digitalizado (imagem). "
                "Tente: 1) Salvar como TXT no seu leitor de PDF, "
                "2) Usar o Google Drive (link) se o PDF original tiver texto, "
                "ou 3) Copiar e colar o conteúdo em um arquivo .txt."
            )
        return text, None

    elif fname.endswith((".docx", ".doc")):
        text = extract_text_from_docx(raw_bytes)
        if not text.strip():
            return "", "Não foi possível extrair texto do documento Word."
        return text, None

    else:
        # TXT, MD, etc.
        try:
            return raw_bytes.decode("utf-8", errors="ignore"), None
        except Exception:
            return "", "Não foi possível ler o arquivo de texto."
